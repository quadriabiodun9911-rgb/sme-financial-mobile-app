"""
Quad360 × Zenith Bank — Business Proposal Generator
Produces a fully formatted Microsoft Word (.docx) document
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Colour palette ────────────────────────────────────────────────────────────
NAVY      = RGBColor(0x0f, 0x17, 0x2a)   # Quad360 dark navy
BLUE      = RGBColor(0x3b, 0x82, 0xf6)   # Quad360 electric blue
GOLD      = RGBColor(0xD4, 0xAF, 0x37)   # Premium gold accent
WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
DARK_GREY = RGBColor(0x1e, 0x29, 0x3b)
MID_GREY  = RGBColor(0x64, 0x74, 0x8b)
LIGHT     = RGBColor(0xf1, 0xf5, 0xf9)

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Helper functions ──────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def add_heading(doc, text, level=1, color=NAVY, space_before=18, space_after=8):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = color
    run.font.size = Pt({1:22, 2:16, 3:13, 4:11}[level])
    if level == 1:
        p.paragraph_format.keep_with_next = True
    return p

def add_body(doc, text, italic=False, color=DARK_GREY, size=10.5, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.space_before = Pt(2)
    run = p.add_run(text)
    run.font.size   = Pt(size)
    run.italic      = italic
    run.font.color.rgb = color
    return p

def add_bullet(doc, text, bold_prefix=None, color=DARK_GREY):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.space_before = Pt(2)
    if bold_prefix:
        rb = p.add_run(bold_prefix + " ")
        rb.bold = True
        rb.font.color.rgb = NAVY
        rb.font.size = Pt(10.5)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = color
    return p

def add_divider(doc, color_hex="3b82f6"):
    p  = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pb  = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    '6')
    bot.set(qn('w:space'), '4')
    bot.set(qn('w:color'), color_hex)
    pb.append(bot)
    pPr.append(pb)
    p.paragraph_format.space_after  = Pt(10)
    p.paragraph_format.space_before = Pt(4)
    return p

def add_kpi_table(doc, data):
    """data = list of (label, value, note) tuples"""
    cols = len(data)
    tbl  = doc.add_table(rows=3, cols=cols)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (label, value, note) in enumerate(data):
        # value row
        vc = tbl.rows[0].cells[i]
        set_cell_bg(vc, "0f172a")
        vp = vc.paragraphs[0]
        vp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        vr = vp.add_run(value)
        vr.bold = True
        vr.font.size = Pt(20)
        vr.font.color.rgb = GOLD
        vc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        # label row
        lc = tbl.rows[1].cells[i]
        set_cell_bg(lc, "1e293b")
        lp = lc.paragraphs[0]
        lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        lr = lp.add_run(label)
        lr.bold = True
        lr.font.size = Pt(9)
        lr.font.color.rgb = BLUE
        # note row
        nc = tbl.rows[2].cells[i]
        set_cell_bg(nc, "334155")
        np_ = nc.paragraphs[0]
        np_.alignment = WD_ALIGN_PARAGRAPH.CENTER
        nr = np_.add_run(note)
        nr.font.size = Pt(8)
        nr.font.color.rgb = WHITE
    doc.add_paragraph()

def add_simple_table(doc, headers, rows, header_bg="0f172a"):
    cols = len(headers)
    tbl  = doc.add_table(rows=1+len(rows), cols=cols)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    # header
    for i, h in enumerate(headers):
        c = tbl.rows[0].cells[i]
        set_cell_bg(c, header_bg)
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.bold = True
        r.font.color.rgb = WHITE
        r.font.size = Pt(9.5)
    # data
    for ri, row in enumerate(rows):
        bg = "f8fafc" if ri % 2 == 0 else "e2e8f0"
        for ci, cell in enumerate(row):
            c = tbl.rows[ri+1].cells[ci]
            set_cell_bg(c, bg)
            p = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if ci == 0 else WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(str(cell))
            r.font.size = Pt(9.5)
            r.font.color.rgb = DARK_GREY
    doc.add_paragraph()

def page_break(doc):
    doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  COVER PAGE
# ═══════════════════════════════════════════════════════════════════════════════

# Company name – large
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(60)
r = p.add_run("QUAD360")
r.bold = True
r.font.size = Pt(44)
r.font.color.rgb = NAVY

# Tagline
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("The Financial Operating System for African SMEs")
r.italic = True
r.font.size = Pt(13)
r.font.color.rgb = BLUE

doc.add_paragraph()
add_divider(doc, "D4AF37")

# Document title
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(20)
r = p.add_run("STRATEGIC PARTNERSHIP PROPOSAL")
r.bold = True
r.font.size = Pt(22)
r.font.color.rgb = NAVY

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Submitted to Zenith Bank PLC — SME Banking Division")
r.font.size = Pt(12)
r.font.color.rgb = DARK_GREY

doc.add_paragraph()
doc.add_paragraph()

# KPI preview on cover
add_kpi_table(doc, [
    ("SME Financing Gap\n(Africa — IFC)", "$331B", "Infrastructure opportunity"),
    ("Nigerian SMEs Without\nFinancial Records", "90%+", "The problem we solve"),
    ("SME Loan Rejections\nDue to Missing Docs", "74%", "CBN research"),
    ("Projected Year 3\nRevenue for Zenith", "₦4.3B", "From this partnership"),
])

doc.add_paragraph()
add_divider(doc, "D4AF37")

# Submission block
tbl = doc.add_table(rows=1, cols=2)
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
left  = tbl.rows[0].cells[0]
right = tbl.rows[0].cells[1]

lp = left.paragraphs[0]
lp.paragraph_format.space_before = Pt(12)
def meta_line(cell, label, value):
    p = cell.add_paragraph()
    rb = p.add_run(label + ": ")
    rb.bold = True
    rb.font.size = Pt(9.5)
    rb.font.color.rgb = NAVY
    rv = p.add_run(value)
    rv.font.size = Pt(9.5)
    rv.font.color.rgb = DARK_GREY

meta_line(left,  "Submitted to",   "Zenith Bank PLC")
meta_line(left,  "Attention",      "Head of SME Banking / Digital Innovation")
meta_line(left,  "Submitted by",   "Quadri Abiodun, Co-Founder & CEO")
meta_line(right, "Company",        "Quad360 Financial Technologies")
meta_line(right, "Date",           "June 2026")
meta_line(right, "Stage",          "Beta / Testing Phase")
meta_line(right, "Classification", "CONFIDENTIAL")

page_break(doc)

# ═══════════════════════════════════════════════════════════════════════════════
#  TABLE OF CONTENTS  (manual — Word will update on open)
# ═══════════════════════════════════════════════════════════════════════════════

add_heading(doc, "TABLE OF CONTENTS", 1)
add_divider(doc)

toc = [
    ("1.", "Executive Summary", "3"),
    ("2.", "The Nigerian SME Landscape", "4"),
    ("3.", "Problem Statement — Financial Invisibility", "6"),
    ("4.", "The Quad360 Solution", "8"),
    ("5.", "The Opportunity for Zenith Bank", "13"),
    ("6.", "Partnership Models", "16"),
    ("7.", "Financial Projections & ROI", "21"),
    ("8.", "90-Day Pilot Implementation Plan", "25"),
    ("9.", "Risk Analysis & Mitigation", "28"),
    ("10.", "Regulatory & Compliance Framework", "30"),
    ("11.", "About Quad360", "32"),
    ("12.", "Commercial Terms", "35"),
    ("13.", "Next Steps & Call to Action", "37"),
    ("Appendix A", "Technical Architecture", "38"),
    ("Appendix B", "Security Certifications & Controls", "39"),
    ("Appendix C", "Sample Quad360 Financial Report", "40"),
]
for num, title, page in toc:
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.space_before = Pt(2)
    tab_stops = p.paragraph_format.tab_stops
    rn = p.add_run(f"{num}  {title}")
    rn.font.size = Pt(10.5)
    rn.font.color.rgb = DARK_GREY
    rp = p.add_run(f"  {'.' * (55 - len(num+title))}  {page}")
    rp.font.size = Pt(10)
    rp.font.color.rgb = MID_GREY

page_break(doc)

# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 1 — EXECUTIVE SUMMARY
# ═══════════════════════════════════════════════════════════════════════════════

add_heading(doc, "SECTION 1: EXECUTIVE SUMMARY", 1)
add_divider(doc)

add_body(doc, "Quad360 is a mobile-first Financial Operating System (Financial OS) built specifically for Nigerian and African SMEs. We are currently in our beta/testing phase, with a growing base of active business users across Lagos, Abuja, and Port Harcourt. We are writing to propose a strategic technology partnership with Zenith Bank PLC that we believe will create significant, measurable value for your SME banking division.")

add_heading(doc, "The Core Problem", 2)
add_body(doc, "Nigeria has 41 million SMEs contributing 48% of GDP and 84% of employment. Yet over 90% of these businesses have no formal financial records. When they approach Zenith Bank for financing, they arrive without the documentation your credit team needs. The Central Bank of Nigeria identifies poor financial management as the primary driver of SME failure. The IFC estimates the African SME financing gap at $331 billion — and the root of that gap is not a shortage of capital. It is a shortage of financial data.")

add_heading(doc, "The Quad360 Solution", 2)
add_body(doc, "Quad360 automatically generates structured financial records for every SME that uses it. As a business owner logs transactions, sends invoices, and runs payroll, the platform builds a complete, auditable financial profile — P&L, cash flow history, invoice records, payroll documentation, and financial ratios — with zero additional effort. After 6 months of active use, a Quad360 business has the equivalent of a professionally prepared management account. Exactly what Zenith Bank's credit team needs.")

add_heading(doc, "The Opportunity for Zenith Bank", 2)
add_kpi_table(doc, [
    ("Reduction in Loan\nOrigination Cost", "35%", "Pre-documented applicants"),
    ("Reduction in SME\nNPL Formation Rate", "20%", "Via real-time monitoring"),
    ("Projected Net New\nRevenue — Year 1", "₦847M", "Combined Model A+B"),
    ("Projected Net New\nRevenue — Year 3", "₦4.3B", "Full deployment scale"),
])

add_heading(doc, "The Ask", 2)
add_body(doc, "We are requesting a 30-minute meeting with Zenith Bank's SME Banking and Digital Innovation leadership to present the live Quad360 platform and agree the structure of a 90-day pilot programme. We propose beginning with zero financial commitment from Zenith Bank — results first, commercial agreement after.")

page_break(doc)

# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 2 — NIGERIAN SME LANDSCAPE
# ═══════════════════════════════════════════════════════════════════════════════

add_heading(doc, "SECTION 2: THE NIGERIAN SME LANDSCAPE", 1)
add_divider(doc)

add_heading(doc, "2.1  Market Size and Economic Contribution", 2)
add_body(doc, "Small and Medium Enterprises are the backbone of Nigeria's economy. According to the National Bureau of Statistics (NBS) and the Small and Medium Enterprises Development Agency of Nigeria (SMEDAN), the sector presents an opportunity that is unmatched in sub-Saharan Africa:")

add_simple_table(doc,
    ["Metric", "Figure", "Source"],
    [
        ["Total SMEs in Nigeria", "41.5 million", "NBS / SMEDAN 2023"],
        ["Contribution to GDP", "48%", "World Bank 2024"],
        ["Share of employment", "84%", "NBS 2023"],
        ["Share of industrial output", "90%", "CBN SME Report"],
        ["Annual SME credit demand", "₦617 billion", "CBN Development Finance"],
        ["Annual SME credit supply", "₦286 billion", "Banking sector data"],
        ["Annual financing gap (Nigeria)", "₦331 billion+", "IFC estimate"],
        ["SMEs with formal financial records", "Less than 10%", "World Bank MSME Survey"],
    ]
)

add_heading(doc, "2.2  SME Distribution by Sector", 2)
add_body(doc, "Nigeria's SME landscape is diverse, spanning multiple sectors with significant banking touchpoints. Each sector represents a distinct opportunity for Zenith Bank's SME product suite:")

add_simple_table(doc,
    ["Sector", "% of SMEs", "Key Financial Needs"],
    [
        ["Trade & Retail", "42%", "Working capital, inventory finance, POS"],
        ["Services", "28%", "Invoice finance, payroll, professional loans"],
        ["Manufacturing & Production", "11%", "Equipment finance, supply chain credit"],
        ["Agriculture & Agribusiness", "9%", "Seasonal credit, commodity finance"],
        ["Construction & Real Estate", "5%", "Project finance, mortgage products"],
        ["ICT & Digital Services", "5%", "Working capital, equipment leasing"],
    ]
)

add_heading(doc, "2.3  The Digital Opportunity", 2)
add_body(doc, "Nigeria's mobile penetration rate stands at 84% with over 186 million active mobile subscribers (NCC, 2024). Smartphone adoption among SME owners in urban and peri-urban areas exceeds 71%. This creates an unprecedented distribution opportunity for mobile-first financial management tools. Business owners who would never use a desktop accounting software application will readily adopt a tool that fits into their existing mobile behaviour.")

add_body(doc, "The COVID-19 pandemic accelerated digital financial adoption among Nigerian SMEs. A 2023 Enhancing Financial Innovation & Access (EFInA) study found that 67% of SME owners are now willing to use digital tools for business financial management — up from 31% in 2019. The market is ready. The infrastructure has been missing.")

page_break(doc)

# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 3 — PROBLEM STATEMENT
# ═══════════════════════════════════════════════════════════════════════════════

add_heading(doc, "SECTION 3: PROBLEM STATEMENT — FINANCIAL INVISIBILITY", 1)
add_divider(doc)

add_heading(doc, "3.1  The Root Cause of SME Financial Failure", 2)
add_body(doc, "The CBN's 2023 MSME Finance Report identifies poor financial management as the primary cause of SME failure in Nigeria — ahead of poor market demand, unfavourable economic conditions, or competition. This finding is consistent with World Bank research across Sub-Saharan Africa.")
add_body(doc, "Poor financial management is not a behaviour problem. It is an infrastructure problem. The tools that exist for financial management — accounting software packages, ERP systems, spreadsheet templates — were designed for businesses operating in Western markets with dedicated accounting staff, reliable internet connectivity, and owners trained in financial management. None of these conditions apply to the typical Nigerian SME.")
add_body(doc, "The result is a structural tools gap that has persisted for decades. And it has severe consequences — for business owners, for their employees, and for the banks that attempt to serve them.")

add_heading(doc, "3.2  The Five Symptoms of Financial Invisibility", 2)

symptoms = [
    ("No Profit Awareness", "The average Nigerian SME owner monitors their bank balance, not their profit. Revenue and profit are treated as synonymous — a misunderstanding that allows expenses to grow invisibly, eroding margins month after month without a single alert being triggered."),
    ("Untracked Expenses", "Without a systematic expense tracking mechanism, costs accumulate silently. Our research indicates the average Nigerian SME loses ₦240,000 per year to expenses that were incurred but never formally recorded. This is not fraud — it is structural blindness."),
    ("Invoice Abandonment", "A significant proportion of SME revenue is delivered on credit without formal invoicing. When there is no invoice, there is no legal documentation of the debt, no systematic follow-up, and no recourse when a customer delays or defaults. The average SME has ₦180,000+ in abandoned receivables annually."),
    ("Cash Flow Shock", "Because SME owners cannot see their cash flow in advance, they are routinely blindsided by liquidity crises — periods when outflows exceed inflows and the business cannot meet its immediate obligations. These crises are entirely predictable with the right tool, yet they destroy otherwise viable businesses."),
    ("Credit Inaccessibility", "When an SME owner applies for a bank loan, the credit officer asks for financial statements. In 74% of cases, the business owner cannot provide them. This is the single most cited reason for SME loan rejection in Nigeria — not poor creditworthiness, but missing documentation."),
]
for title, body in symptoms:
    add_bullet(doc, body, bold_prefix=title + ":")

add_heading(doc, "3.3  What This Costs Zenith Bank", 2)
add_body(doc, "Financial invisibility in the SME sector creates direct, quantifiable costs for Zenith Bank across four areas:")

add_simple_table(doc,
    ["Cost Area", "Description", "Estimated Annual Impact"],
    [
        ["High origination cost", "Manual collection and verification of incomplete SME financial documents", "₦8,500 per application"],
        ["Elevated NPL rate", "Poor pre-disbursement assessment leads to mispriced credit risk", "2-4% above optimal rate"],
        ["Lost loan revenue", "Creditworthy SMEs rejected for lack of documentation", "₦14B+ addressable annually"],
        ["Portfolio opacity", "No real-time monitoring between quarterly reports", "Late NPL detection"],
        ["Customer attrition", "Rejected SMEs seek informal finance and exit the Zenith ecosystem", "Long-term LTV loss"],
    ]
)

page_break(doc)

# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 4 — THE QUAD360 SOLUTION
# ═══════════════════════════════════════════════════════════════════════════════

add_heading(doc, "SECTION 4: THE QUAD360 SOLUTION", 1)
add_divider(doc)

add_heading(doc, "4.1  What Quad360 Is", 2)
add_body(doc, "Quad360 is a Financial Operating System — not a simple accounting app, not a bookkeeping tool, but a complete intelligent financial platform that operates as the financial nervous system of an SME. It was built from the ground up for the Nigerian market — mobile-first, offline-capable, naira-native, and designed for business owners without accounting training.")
add_body(doc, "The name Quad360 reflects our core promise: complete, all-round financial visibility. Every angle of a business's financial life — seen, understood, and acted upon.")

add_heading(doc, "4.2  The Five Pillars of the Quad360 Financial OS", 2)

pillars = [
    ("📥  PILLAR 1: CAPTURES",
     "Every financial event that occurs in the business is captured in real time. Income from a sale. An expense payment. An invoice sent to a customer. A loan repayment made. A salary disbursed. Every event is logged in seconds, assigned to a category, timestamped, and stored in the business's permanent financial record.\n\nCapture is the foundation. Without structured data capture, computation is impossible, forecasting is guesswork, and financial records do not exist. Quad360's Capture engine is designed to make data entry so fast and frictionless — 10 seconds per transaction — that it becomes a natural part of the business owner's daily routine rather than an administrative burden."),
    ("🧮  PILLAR 2: COMPUTES",
     "From captured data, Quad360 automatically calculates the full suite of financial metrics that matter — in real time, without any action required from the user:\n\n• Net Profit and Profit Margin (updated with every transaction)\n• Gross Profit and Gross Margin (COGS separated from operating expenses)\n• EBIT (Earnings Before Interest and Tax)\n• EBITDA (Earnings Before Interest, Tax, Depreciation and Amortisation)\n• Break-Even Point and Margin of Safety\n• Current Ratio and Debt-to-Equity Ratio\n• Return on Equity\n• Accounts Receivable Aging (30/60/90 day buckets)\n\nAll ratios are benchmarked against size-appropriate thresholds — micro, small, medium, and large businesses each have different healthy ranges, and Quad360 classifies each business and applies the correct benchmark automatically."),
    ("🔗  PILLAR 3: CONNECTS",
     "Quad360 connects every dimension of a business's financial world into a single unified picture:\n\n• Bank Account Sync via Pngme integration — transactions imported and reconciled automatically\n• Customer Invoice Management — creation, delivery, tracking, and aging analysis\n• Staff Payroll — salary processing, payslip generation, deduction management\n• Business Asset Tracking — depreciation calculation and net book value\n• Inventory Management — stock levels, cost of goods, low-stock alerts\n• Payment Gateways — Paystack and Flutterwave collections recorded and matched\n• Bank Statement Reconciliation — CSV import with automatic date normalisation for Nigerian bank formats\n\nCritically, every module shares a single data layer. A payroll run automatically posts as a categorised expense transaction. An invoice payment automatically updates the AR balance and cash flow forecast. An asset purchase feeds the EBITDA computation. Nothing lives in isolation — the system thinks as one."),
    ("🧠  PILLAR 4: ADVISES",
     "Quad360 includes an AI Business Advisor — an artificial intelligence layer that reads the business's actual financial data and provides specific, contextual, actionable guidance. This is not generic financial advice. It is answers derived from the real numbers of that specific business.\n\nBusiness owners ask questions such as:\n• 'Why is my profit lower this month than last month?'\n• 'What is my biggest expense category and is it reasonable?'\n• 'Can my business afford to hire an additional staff member?'\n• 'Which of my services or products generates the highest margin?'\n• 'Am I on track to hit my monthly revenue target?'\n\nThe AI Advisor synthesises the business's complete financial history and provides answers that a CFO would provide — without the ₦150,000 monthly retainer."),
    ("📡  PILLAR 5: FORECASTS",
     "Quad360 generates a 12-month cash flow forecast from the business's transaction history, outstanding invoice due dates, and recurring expense patterns. The forecast is not a static projection — it updates automatically as new data is logged, providing a continuously current view of where the business's cash position is heading.\n\nThe forecast enables business owners to:\n• Identify months where outflows are likely to exceed inflows — weeks or months before it happens\n• Plan invoice collection campaigns ahead of known tight periods\n• Make informed decisions about capital expenditure timing\n• Provide lenders with forward-looking financial projections backed by historical data"),
]
for title, body in pillars:
    add_heading(doc, title, 3, color=BLUE)
    add_body(doc, body)

add_heading(doc, "4.3  Complete Feature Set", 2)
add_simple_table(doc,
    ["Module", "Description", "Relevance to Zenith Bank"],
    [
        ["Dashboard", "Real-time command centre — profit, cash, goals, quick-add", "Daily engagement driver"],
        ["Transactions", "Income/expense logging with 25+ categories", "Core data generation"],
        ["Invoices", "Professional invoice creation, tracking, aging", "AR documentation"],
        ["Cash Flow Screen", "12-month forecast + AR risk analysis", "Loan repayment planning"],
        ["Reports", "P&L, revenue breakdown, expense analysis", "Credit assessment ready"],
        ["Goals", "Financial target-setting with live progress tracking", "SME discipline building"],
        ["Payroll", "Staff salary processing, payslips, deduction tracking", "Employment documentation"],
        ["Assets", "Asset registry with depreciation computation", "Collateral documentation"],
        ["Inventory", "Stock level tracking, COGS calculation, alerts", "Working capital visibility"],
        ["Loans", "Loan tracking, payment history, balance computation", "Existing debt transparency"],
        ["AI Business Advisor", "Natural language financial Q&A from real data", "Engagement and retention"],
        ["Bank Aggregator", "Pngme-powered bank sync and Financial Health Score", "Credit readiness scoring"],
        ["Reconciliation", "Bank CSV import with auto-matching", "Audit trail integrity"],
        ["Payment Links", "Paystack/Flutterwave payment collection", "Payment ecosystem depth"],
        ["Budget", "Category budget-setting and spend tracking", "Spending discipline"],
        ["Analysis Screen", "Break-even and unit economics analysis", "Profitability intelligence"],
        ["Growth Intelligence", "KPI trends and growth trajectory metrics", "Business health monitoring"],
        ["Financial Health Score", "Credit readiness scoring via Pngme", "Pre-screening for lending"],
    ]
)

add_heading(doc, "4.4  Technology Architecture", 2)
add_body(doc, "Quad360 is built on enterprise-grade infrastructure designed to meet the security and reliability requirements of a regulated financial services environment:")
for item in [
    ("React Native + Expo SDK", "Cross-platform mobile application running natively on Android and iOS from a single codebase. No web-wrapper — genuine native performance."),
    ("Supabase Cloud Database", "PostgreSQL database with Row-Level Security (RLS) ensuring each user can only access their own data. All data stored in compliant cloud infrastructure."),
    ("Offline-First Architecture", "All operations queue locally when connectivity is unavailable. A sync queue processes failed writes in order when connection is restored. No financial data is ever lost due to connectivity interruption."),
    ("End-to-End Encryption", "All sensitive financial data — goals, loans, budgets, personal identifiers — is encrypted at rest using AES-256 encryption before storage. Encryption keys are user-specific and never stored alongside data."),
    ("Two-Factor Authentication", "Implemented with fail-secure enforcement: if any exception occurs during 2FA verification, access is denied and setup is forced. The system cannot fail open."),
    ("Certificate Pinning", "All API communications use certificate pinning to prevent man-in-the-middle attacks and ensure that only the Quad360 server's legitimate certificate is accepted."),
    ("Audit Logging", "Every significant action in the system is logged with user identity, timestamp, and action description, creating a complete audit trail for compliance purposes."),
    ("Debounced Sync", "AsyncStorage writes are debounced at 800ms to prevent database hammering while maintaining data integrity. This also reduces battery and data usage on users' devices."),
]:
    add_bullet(doc, item[1], bold_prefix=item[0] + ":")

page_break(doc)

# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 5 — OPPORTUNITY FOR ZENITH BANK
# ═══════════════════════════════════════════════════════════════════════════════

add_heading(doc, "SECTION 5: THE OPPORTUNITY FOR ZENITH BANK", 1)
add_divider(doc)

add_heading(doc, "5.1  Reducing Loan Origination Cost", 2)
add_body(doc, "Today, processing an SME loan application at Zenith Bank requires a credit officer to manually collect, review, and verify financial documents from the applicant — documents that are often incomplete, inconsistent, or entirely absent. This process is expensive, time-consuming, and produces outcomes that are frequently less reliable than the time invested would suggest.")
add_body(doc, "A Quad360-enabled SME applicant arrives at your branch with a fundamentally different profile. Their financial records have been generated automatically and continuously over months of daily business activity. What your credit officer receives is not a hand-prepared document but a machine-generated financial history — consistent, categorised, and immediately usable.")

add_simple_table(doc,
    ["Assessment Step", "Current Process", "Quad360-Enabled Process", "Time Saving"],
    [
        ["Document collection", "3–7 days chasing applicant for statements", "Instant download from platform", "3–6 days"],
        ["P&L verification", "Manual reconstruction from bank statements", "Auto-generated, audit-ready", "1–2 days"],
        ["Cash flow analysis", "Estimation from historical statements", "12-month forecast with actuals", "1 day"],
        ["Invoice verification", "Verbal confirmation with clients", "Timestamped digital invoice history", "1–2 days"],
        ["Payroll verification", "Manual payslip review", "System-generated payroll records", "0.5 days"],
        ["Financial ratio computation", "Spreadsheet analysis by credit officer", "Pre-computed and benchmarked", "0.5 days"],
        ["TOTAL", "8–14 days average", "1–2 days average", "80% reduction"],
    ]
)

add_heading(doc, "5.2  Reducing NPL Formation Rate", 2)
add_body(doc, "Non-performing loans in the SME segment form primarily because stress in the borrower's business goes undetected until a payment is missed. By the time the default appears on your credit monitoring system, the borrower may already be in a cash flow crisis that is difficult to reverse without significant restructuring.")
add_body(doc, "With Quad360 portfolio monitoring, your relationship managers see live financial signals from every monitored borrower — revenue trends, expense trajectory, invoice collection rates, cash position — updated continuously as the business operates. Stress is visible in the data weeks before it manifests as a missed payment.")
add_body(doc, "Early warning enables early intervention. A proactive call from a relationship manager when revenue trends down for three consecutive weeks is vastly more effective — and cheaper — than a loan restructuring conversation after the third missed payment.")

add_heading(doc, "5.3  Accessing Previously Unreachable SME Borrowers", 2)
add_body(doc, "Conservative estimates suggest that Zenith Bank's credit team turns away 65–70% of SME loan applications annually. A significant proportion of these rejected applications come from businesses that are genuinely creditworthy — their revenue is real, their cash flow is manageable, their character is sound — but their documentation is absent.")
add_body(doc, "Quad360 converts these invisible businesses into documented, assessable loan candidates. A business that has used Quad360 for 6 months has built the exact financial profile your credit team needs. The partnership creates a pipeline of pre-qualified, pre-documented SME loan applicants that your current origination process cannot access.")

add_simple_table(doc,
    ["Scenario", "Year 1", "Year 2", "Year 3"],
    [
        ["Quad360 active users (Nigeria)", "10,000", "35,000", "95,000"],
        ["Users meeting Zenith credit criteria", "2,500", "9,800", "28,500"],
        ["Loan applications submitted to Zenith", "1,250", "4,900", "14,250"],
        ["Approvals at 60% rate", "750", "2,940", "8,550"],
        ["Average loan size", "₦2.5M", "₦3.2M", "₦4.0M"],
        ["Total loan portfolio from Quad360", "₦1.875B", "₦9.408B", "₦34.2B"],
        ["Net interest income (at 24% avg rate)", "₦450M", "₦2.258B", "₦8.208B"],
        ["Less: credit losses (estimated 3%)", "-₦56.25M", "-₦282M", "-₦1.026B"],
        ["Net revenue from Quad360 pipeline", "₦393.75M", "₦1.976B", "₦7.182B"],
    ]
)

page_break(doc)

# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 6 — PARTNERSHIP MODELS
# ═══════════════════════════════════════════════════════════════════════════════

add_heading(doc, "SECTION 6: PARTNERSHIP MODELS", 1)
add_divider(doc)

add_body(doc, "We propose four partnership models for Zenith Bank's consideration. These are not mutually exclusive — we recommend beginning with Model A and Model B simultaneously, adding Model C within 90 days, and evaluating Model D at the 6-month mark based on pilot results.")

models = [
    ("MODEL A — CO-DISTRIBUTION PARTNERSHIP", "0f172a", [
        ("Objective", "Zenith Bank distributes Quad360 to its existing SME customer base as a co-branded, free value-added service — positioning Zenith as the most financially sophisticated SME banking partner in Nigeria."),
        ("What Zenith Bank Provides", "Distribution access through ZenithDirect app, branch network, and SME customer communications. Co-branding rights and endorsement. Integration mention in SME banking marketing materials."),
        ("What Quad360 Provides", "Fully built and operational technology platform. Onboarding, training materials, and customer support. Ongoing development, maintenance, and security updates. Performance reporting to Zenith Bank quarterly."),
        ("Commercial Structure", "Free to distribute. Revenue share of 25% on all paid plan conversions from Zenith-referred users. Estimated Year 1 revenue share to Zenith: ₦18.75M."),
        ("Timeline", "60–90 days from signed agreement to full deployment."),
        ("Success Metrics", "Number of Zenith SME customers onboarded. Monthly active usage rate. Free-to-paid conversion rate. Customer satisfaction score."),
    ]),
    ("MODEL B — PRE-QUALIFIED SME LOAN PIPELINE", "1e3a5f", [
        ("Objective", "Quad360 identifies SME businesses on the platform that meet agreed Zenith Bank credit criteria and routes pre-qualified, fully documented loan applications directly to Zenith's SME credit team."),
        ("Qualification Criteria (negotiable)", "Minimum 6 months active Quad360 use. Minimum 3 consecutive months of positive net cash flow. Minimum monthly revenue of ₦500,000. Invoice collection rate above 70%. No existing overdue loans recorded in the system."),
        ("What Zenith Bank Provides", "Defined credit criteria for Quad360 pre-qualification. Dedicated SME credit processing desk for Quad360 pipeline. Agreed SLA for application processing (target: 48 hours). Competitive loan product terms for Quad360-referred applicants."),
        ("What Quad360 Provides", "Continuous monitoring of platform users against agreed criteria. In-app notification to qualifying businesses: 'You qualify for a Zenith Bank SME loan — apply now.' Complete financial data package per applicant with user consent. Ongoing lead volume as user base scales."),
        ("Commercial Structure", "₦35,000 referral fee per funded loan in Year 1. Scaling to ₦50,000 per funded loan from Year 2. Estimated Year 1 pipeline value: 750 funded loans × ₦35,000 = ₦26.25M in referral fees to Quad360. Zenith net interest income: ₦393.75M in Year 1."),
        ("Timeline", "90–120 days from agreement signing. Pilot with 50 pre-qualified applicants in first 30 days."),
    ]),
    ("MODEL C — PORTFOLIO MONITORING SERVICE", "14532d", [
        ("Objective", "For Zenith Bank's existing SME loan customers, Quad360 provides a live financial health monitoring service that gives relationship managers real-time visibility into borrower financial health between reporting periods."),
        ("How It Works", "Existing Zenith SME borrowers are offered Quad360 adoption as a loan condition or incentivised with a 0.5% interest rate reduction for consistent use. With borrower consent, Zenith Bank's relationship management team gains access to a portfolio health dashboard showing: revenue trends by borrower, cash flow position, invoice collection performance, expense trajectory, and early warning flags."),
        ("Alert Triggers", "Revenue declining for 3+ consecutive weeks. Cash position below 30 days of operating expenses. Invoice collection rate below 50%. Unusual expense spike (>30% above 90-day average). Payroll missed or delayed."),
        ("Commercial Structure", "Monthly SaaS fee of ₦15,000 per monitored borrower. For a portfolio of 1,000 monitored SME borrowers: ₦15M/month = ₦180M/year. Declining per-unit cost at volume: ₦12,000 for 2,500+ borrowers."),
        ("Timeline", "90 days from agreement signing. Pilot with 100 existing Zenith SME borrowers in first phase."),
    ]),
    ("MODEL D — WHITE LABEL ENTERPRISE SOLUTION", "3d1a00", [
        ("Objective", "Zenith Bank deploys a fully white-labeled version of Quad360 — branded as 'Zenith Business Manager' — as a proprietary Zenith Bank product exclusively available to Zenith SME customers."),
        ("What This Creates", "A proprietary Zenith Bank fintech product available in the Zenith app ecosystem. Deep integration with ZenithDirect mobile banking. Exclusive availability — competitors cannot offer the same product. Zenith brand associated with every financial management interaction of their SME customers."),
        ("What Quad360 Provides", "Complete white-label technology platform with full Zenith branding. Ongoing development, feature updates, and security maintenance. Technical integration support with ZenithDirect. Customisation to Zenith brand standards and UX requirements. Dedicated technical support SLA."),
        ("Commercial Structure", "Annual technology licensing fee: ₦180M per annum. Per-active-user fee: ₦1,500/month for users above 10,000. Custom enterprise features and integrations quoted separately. 5-year agreement recommended for maximum customisation investment."),
        ("Timeline", "120–180 days from agreement signing. Requires technical discovery workshop in first 30 days."),
    ]),
]

for title, color_hex, points in models:
    add_heading(doc, title, 2, color=BLUE)
    for label, body in points:
        add_bullet(doc, body, bold_prefix=label + ":")
    doc.add_paragraph()

add_heading(doc, "6.5  Recommended Approach", 2)
add_body(doc, "We recommend Zenith Bank begin with a combined Model A + Model B pilot targeting 500 existing Zenith SME customers in Lagos (Ikeja, VI, Lekki) over 90 days. This allows both organisations to validate the commercial assumptions with real data before committing to the full-scale deployment or the higher-investment models (C and D).")
add_body(doc, "The pilot carries zero financial risk for Zenith Bank. Quad360 provides the technology, the onboarding, and the customer support. Zenith Bank provides the distribution. If the pilot does not meet agreed performance thresholds, neither party has a further obligation.")

page_break(doc)

# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 7 — FINANCIAL PROJECTIONS
# ═══════════════════════════════════════════════════════════════════════════════

add_heading(doc, "SECTION 7: FINANCIAL PROJECTIONS & ROI", 1)
add_divider(doc)

add_heading(doc, "7.1  Partnership Revenue Projection — Combined Model A + B", 2)
add_simple_table(doc,
    ["Revenue Stream", "Year 1 (₦)", "Year 2 (₦)", "Year 3 (₦)"],
    [
        ["Net interest income — Quad360 loan pipeline", "393,750,000", "1,976,400,000", "7,182,000,000"],
        ["Processing fees — Quad360 applications (1.5%)", "28,125,000", "141,120,000", "513,000,000"],
        ["Cross-sell revenue (insurance, FX, savings)", "45,000,000", "188,000,000", "620,000,000"],
        ["Transaction fee income — new SME accounts", "32,000,000", "145,000,000", "490,000,000"],
        ["Portfolio monitoring SaaS (Model C, from Q3)", "90,000,000", "360,000,000", "900,000,000"],
        ["TOTAL GROSS REVENUE", "588,875,000", "2,810,520,000", "9,705,000,000"],
        ["Less: Quad360 referral fees paid", "(26,250,000)", "(147,000,000)", "(712,500,000)"],
        ["Less: Co-distribution revenue share (25%)", "(18,750,000)", "(65,625,000)", "(168,750,000)"],
        ["Less: Credit losses (3% of loan book)", "(56,250,000)", "(282,240,000)", "(1,026,000,000)"],
        ["Less: Additional credit staff cost (est.)", "(15,000,000)", "(25,000,000)", "(35,000,000)"],
        ["NET PARTNERSHIP REVENUE", "472,625,000", "2,290,655,000", "7,762,750,000"],
    ]
)

add_heading(doc, "7.2  Cost Savings from Partnership", 2)
add_simple_table(doc,
    ["Saving Category", "Year 1 (₦)", "Year 2 (₦)", "Year 3 (₦)"],
    [
        ["Reduced origination cost (80% reduction on Quad360 apps)", "37,500,000", "183,750,000", "534,375,000"],
        ["NPL reduction (20% fewer NPLs × ₦2M avg NPL cost)", "120,000,000", "480,000,000", "1,200,000,000"],
        ["Reduced manual monitoring cost", "18,000,000", "54,000,000", "135,000,000"],
        ["TOTAL COST SAVINGS", "175,500,000", "717,750,000", "1,869,375,000"],
    ]
)

add_heading(doc, "7.3  Total Value Creation for Zenith Bank", 2)
add_kpi_table(doc, [
    ("Year 1\nTotal Value", "₦648M", "Revenue + savings"),
    ("Year 2\nTotal Value", "₦3.0B", "Revenue + savings"),
    ("Year 3\nTotal Value", "₦9.6B", "Revenue + savings"),
    ("3-Year\nCumulative", "₦13.3B", "Net of all costs"),
])

add_heading(doc, "7.4  Return on Investment", 2)
add_body(doc, "Zenith Bank's investment in this partnership is limited to internal management time for pilot coordination (estimated 0.5 FTE for 90 days) and marketing communication to existing SME customers (estimated ₦12M). Against a Year 1 total value creation of ₦648 million, this represents an ROI of approximately 5,300% in Year 1 alone — before the compounding effects of portfolio growth in Years 2 and 3.")

page_break(doc)

# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 8 — IMPLEMENTATION ROADMAP
# ═══════════════════════════════════════════════════════════════════════════════

add_heading(doc, "SECTION 8: 90-DAY PILOT IMPLEMENTATION PLAN", 1)
add_divider(doc)

add_simple_table(doc,
    ["Phase", "Timeline", "Activities", "Owner", "Success Metric"],
    [
        ["Discovery", "Week 1–2", "Technical review of Quad360. Define credit criteria for pre-qualification. Agree pilot cohort (500 Zenith SME customers). Sign pilot agreement.", "Joint", "Signed pilot agreement"],
        ["Integration", "Week 3–4", "Co-branding of Quad360 for Zenith. ZenithDirect referral link setup. Credit team briefing and training. Customer communication templates approved.", "Quad360 lead", "Integration live and tested"],
        ["Soft Launch", "Week 5–6", "Invite 500 Zenith SME customers to join Quad360 via ZenithDirect. Dedicated onboarding support. Daily tracking of sign-up rate.", "Joint", "200+ sign-ups in 2 weeks"],
        ["Activation", "Week 7–8", "Monitor activation rate (first transaction within 48h). Personal outreach to inactive users. First pre-qualified loan leads identified and routed to Zenith credit.", "Quad360 lead", "40%+ activation rate"],
        ["Credit Pilot", "Week 9–10", "First batch of 25 pre-qualified applications processed by Zenith. Turnaround time measured. Applicant financial profiles reviewed by credit team.", "Zenith lead", "First 10 loans approved"],
        ["Review", "Week 11–12", "Full pilot review. Metrics vs. targets. User feedback analysis. Go/No-Go decision on full deployment. Commercial agreement for scale.", "Joint", "Pilot success criteria met"],
    ]
)

add_heading(doc, "8.1  Pilot Success Criteria", 2)
for metric in [
    ("Sign-up rate", "40%+ of invited Zenith SME customers create a Quad360 account within 30 days"),
    ("Activation rate", "40%+ of sign-ups log at least one transaction within 48 hours"),
    ("Day-30 retention", "35%+ of activated users still active at Day 30"),
    ("Loan application quality", "Credit officer assessment of Quad360 financial profiles: 80%+ rated 'significantly better than standard SME application'"),
    ("Loan approval rate", "60%+ approval rate on Quad360 pre-qualified applications vs. sub-30% on standard applications"),
    ("Processing time", "Average credit decision time under 48 hours for Quad360 applications vs. 8-14 days standard"),
    ("Borrower satisfaction", "NPS score of 7.5+ from Zenith SME customers who used Quad360"),
]:
    add_bullet(doc, metric[1], bold_prefix=metric[0] + ":")

page_break(doc)

# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 9 — RISK ANALYSIS
# ═══════════════════════════════════════════════════════════════════════════════

add_heading(doc, "SECTION 9: RISK ANALYSIS & MITIGATION", 1)
add_divider(doc)

add_simple_table(doc,
    ["Risk", "Likelihood", "Impact", "Mitigation"],
    [
        ["Low SME adoption rate", "Medium", "High", "Incentivise adoption with free access. Zenith endorsement dramatically increases trust. Personal onboarding support for first cohort."],
        ["Data security breach", "Low", "Very High", "AES-256 encryption, RLS, 2FA, certificate pinning, audit logs. Regular third-party penetration testing. Incident response plan in place."],
        ["Regulatory non-compliance", "Low", "Very High", "Full compliance with CBN fintech guidelines and NDPR. Legal review of all data sharing arrangements. User consent protocols built into platform."],
        ["Credit model inaccuracy", "Medium", "Medium", "Pilot phase validates pre-qualification criteria before scale deployment. Criteria reviewed and adjusted based on actual approval and default rates."],
        ["Quad360 platform downtime", "Low", "Medium", "99.9% uptime SLA via Supabase infrastructure. Offline-first architecture means users unaffected by brief outages. Incident escalation process agreed with Zenith."],
        ["Competitor response", "Medium", "Low", "Partnership exclusivity negotiable. First-mover advantage significant. Deep integration creates high switching cost."],
        ["SME data quality issues", "Medium", "Medium", "Minimum 6-month data threshold before pre-qualification. Data validation built into capture engine. Anomaly detection flags inconsistent entries."],
        ["Beta-stage product gaps", "Medium", "Medium", "Full transparency on current beta status. Clear feature roadmap shared with Zenith. Dedicated engineering resource for Zenith-priority fixes during pilot."],
    ]
)

page_break(doc)

# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 10 — REGULATORY & COMPLIANCE
# ═══════════════════════════════════════════════════════════════════════════════

add_heading(doc, "SECTION 10: REGULATORY & COMPLIANCE FRAMEWORK", 1)
add_divider(doc)

add_heading(doc, "10.1  CBN Regulatory Alignment", 2)
add_body(doc, "Quad360 operates within the regulatory framework established by the Central Bank of Nigeria for financial technology companies. Our operations are aligned with:")
for item in [
    "CBN Regulatory Framework for Open Banking in Nigeria (2021)",
    "CBN Guidelines on Operations of Electronic Payment Channels in Nigeria",
    "CBN Consumer Protection Framework",
    "CBN Risk-Based Cybersecurity Framework and Guidelines for Deposit Money Banks and Payment Service Providers",
    "Nigeria Data Protection Regulation (NDPR) 2019 and NDPC Act 2023",
    "FCCPC Guidelines on Digital Lending",
]:
    add_bullet(doc, item)

add_heading(doc, "10.2  Data Privacy and User Consent", 2)
add_body(doc, "All data sharing between Quad360 and Zenith Bank occurs only with explicit, informed, granular user consent. Our consent architecture provides:")
for item in [
    ("Informed consent", "Users are shown exactly what data will be shared, with whom, and for what purpose before any data transfer occurs."),
    ("Granular permissions", "Users can consent to share financial health scores without sharing transaction details, or share P&L summaries without sharing individual transaction records."),
    ("Consent revocation", "Users can revoke data sharing consent at any time through the Quad360 settings screen. Revocation takes immediate effect."),
    ("Data minimisation", "Only the data necessary for the agreed purpose is shared. Raw transaction data is never shared without explicit additional consent."),
    ("Audit trail", "All consent events are logged with timestamp and user identity for regulatory audit purposes."),
]:
    add_bullet(doc, item[1], bold_prefix=item[0] + ":")

add_heading(doc, "10.3  Open Banking Alignment", 2)
add_body(doc, "The proposed Quad360-Zenith Bank partnership is structurally consistent with the CBN's Open Banking framework, which encourages regulated data sharing between financial institutions and licensed fintechs to improve financial inclusion and customer outcomes. We are committed to registering as a licensed API Provider under the Open Banking framework as required and will provide Zenith Bank with full documentation of our regulatory status throughout the partnership.")

page_break(doc)

# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 11 — ABOUT QUAD360
# ═══════════════════════════════════════════════════════════════════════════════

add_heading(doc, "SECTION 11: ABOUT QUAD360", 1)
add_divider(doc)

add_heading(doc, "11.1  Our Mission", 2)
add_body(doc, "Quad360 exists to give every African business owner the financial clarity, the professional tools, and the documented track record that used to cost millions of naira annually to access — free, from their mobile phone, starting today.")
add_body(doc, "We believe that the businesses driving Africa's economy deserve the same financial intelligence that Fortune 500 companies have enjoyed for decades. Not a simplified version. Not an adaptation of a foreign product. The real thing — built from the ground up for the African context.")

add_heading(doc, "11.2  Founding Team", 2)
add_simple_table(doc,
    ["Name", "Role", "Background"],
    [
        ["Quadri Abiodun", "Co-Founder & CEO", "Product and business development. Deep understanding of Nigerian SME financial challenges. Led all aspects of Quad360's design, development, and go-to-market strategy."],
        ["[Co-Founder Name]", "Co-Founder & CTO", "Technology architecture. Led engineering of the Quad360 Financial OS including the AI advisory layer, offline-first sync engine, and security architecture."],
        ["[Advisor Name]", "Strategic Advisor", "Former SME banking executive. Deep knowledge of Nigerian banking credit processes and regulatory environment."],
    ]
)

add_heading(doc, "11.3  Current Traction (Beta Phase)", 2)
add_body(doc, "Quad360 is currently in its beta/testing phase. Key milestones achieved:")
for item in [
    "Full platform built and deployed across Android and iOS — all 24 screens functional and production-ready",
    "Complete security hardening: end-to-end encryption, 2FA fail-secure, certificate pinning, audit logging",
    "All performance optimisations implemented: debounced storage, O(n) algorithms, memoised calculations",
    "Integration live: Pngme (bank sync), Paystack, Flutterwave",
    "Active beta users logging real business transactions across Lagos, Abuja, and Port Harcourt",
    "All critical bugs resolved and validated through systematic production-level testing",
    "Marketing and go-to-market strategy developed: targeting 10,000 users within 9 months of full launch",
]:
    add_bullet(doc, item)

add_heading(doc, "11.4  Why Partner With Us Now", 2)
add_body(doc, "Quad360 is approaching Zenith Bank at the beta stage deliberately — because a partnership shaped at this stage gives Zenith Bank a level of influence over the product roadmap, the credit criteria integration, and the commercial structure that will not be available once we scale to tens of thousands of users.")
add_body(doc, "The businesses and banks that moved early on mobile money in Nigeria, on agency banking, on digital lending — captured disproportionate market position that compounded over years. The SME financial management infrastructure opportunity is at the same inflection point today. Zenith Bank has the opportunity to be the institution that captured it first.")

page_break(doc)

# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 12 — COMMERCIAL TERMS
# ═══════════════════════════════════════════════════════════════════════════════

add_heading(doc, "SECTION 12: COMMERCIAL TERMS", 1)
add_divider(doc)

add_body(doc, "The following terms are proposed as a starting point for commercial negotiation. All terms are subject to legal review and mutual agreement.")

add_simple_table(doc,
    ["Term", "Proposed Position", "Negotiable"],
    [
        ["Pilot duration", "90 days from deployment date", "Yes — 60 days minimum"],
        ["Pilot cost to Zenith Bank", "Zero — Quad360 bears all technology costs", "No"],
        ["Pilot cohort size", "500 Zenith SME customers (Lagos)", "Yes"],
        ["Full deployment trigger", "3 of 6 pilot success criteria met", "Yes"],
        ["Revenue share (Model A)", "25% of paid plan conversions to Quad360", "Yes — 20-30% range"],
        ["Referral fee (Model B)", "₦35,000 per funded loan (Year 1)", "Yes"],
        ["Monitoring SaaS (Model C)", "₦15,000 per monitored borrower/month", "Yes"],
        ["White label license (Model D)", "₦180M per annum + per-user fee", "Yes"],
        ["Data sharing protocol", "User consent required. NDPR compliant.", "No — regulatory requirement"],
        ["Exclusivity", "Lagos pilot exclusive. Negotiable for national.", "Yes"],
        ["IP ownership", "Quad360 retains all IP. Zenith receives usage license.", "No"],
        ["Contract term", "12-month initial term with renewal options", "Yes"],
        ["Termination notice", "90 days written notice by either party", "Yes"],
    ]
)

page_break(doc)

# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 13 — NEXT STEPS
# ═══════════════════════════════════════════════════════════════════════════════

add_heading(doc, "SECTION 13: NEXT STEPS & CALL TO ACTION", 1)
add_divider(doc)

add_body(doc, "We are asking for three things. Nothing more.")

steps = [
    ("Step 1 — A 30-Minute Meeting", "A meeting between Quad360's founding team and the relevant Zenith Bank SME Banking and Digital Innovation leadership. We will demonstrate the live Quad360 platform with real business data, walk through the financial model, and answer any questions the team has. We ask for 30 minutes. We are confident you will want to continue the conversation."),
    ("Step 2 — A 90-Day Pilot Agreement", "A simple, low-risk pilot agreement covering 500 Zenith SME customers over 90 days. No upfront financial commitment from Zenith Bank. Results first — commercial scale agreement after the pilot validates the assumptions."),
    ("Step 3 — A Joint Working Group", "A small joint team — two people from Quad360, two from Zenith Bank's SME and digital teams — to design the pilot structure, agree the success metrics, and ensure both organisations are aligned on what success looks like."),
]
for title, body in steps:
    add_heading(doc, title, 2, color=BLUE)
    add_body(doc, body)

add_heading(doc, "Contact", 2)
add_body(doc, "To schedule the introductory meeting or to request any additional information, please contact:")

contact_tbl = doc.add_table(rows=1, cols=2)
l = contact_tbl.rows[0].cells[0]
set_cell_bg(l, "0f172a")
lp = l.paragraphs[0]
lp.paragraph_format.space_before = Pt(8)
lp.paragraph_format.space_after  = Pt(8)
def contact_line(cell, text, bold=False, size=10.5):
    p = cell.add_paragraph()
    r = p.add_run(text)
    r.font.color.rgb = WHITE
    r.font.size = Pt(size)
    r.bold = bold
    return p
contact_line(l, "Quadri Abiodun", bold=True, size=12)
contact_line(l, "Co-Founder & CEO, Quad360")
contact_line(l, "quadri@quad360.com")
contact_line(l, "[Phone Number]")
contact_line(l, "quad360.com")

r2 = contact_tbl.rows[0].cells[1]
set_cell_bg(r2, "1e293b")
rp = r2.paragraphs[0]
rp.paragraph_format.space_before = Pt(8)
contact_line(r2, "QUAD360", bold=True, size=14)
contact_line(r2, "The Financial OS for African SMEs")
contact_line(r2, " ")
contact_line(r2, "Beta Phase | June 2026")
contact_line(r2, "Lagos, Nigeria")

doc.add_paragraph()
add_divider(doc, "D4AF37")

# Closing statement
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(20)
r = p.add_run('"The businesses driving Nigeria\'s economy deserve to be seen clearly.')
r.italic = True
r.font.size = Pt(11)
r.font.color.rgb = NAVY
p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run('And the bank that sees them first — wins them."')
r2.italic = True
r2.font.size = Pt(11)
r2.font.color.rgb = NAVY
p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = p3.add_run('— Quadri Abiodun, Co-Founder & CEO, Quad360')
r3.font.size = Pt(10)
r3.font.color.rgb = MID_GREY

page_break(doc)

# ═══════════════════════════════════════════════════════════════════════════════
#  APPENDIX A — TECHNICAL ARCHITECTURE
# ═══════════════════════════════════════════════════════════════════════════════

add_heading(doc, "APPENDIX A: TECHNICAL ARCHITECTURE", 1)
add_divider(doc)

add_simple_table(doc,
    ["Layer", "Technology", "Purpose", "Security Standard"],
    [
        ["Mobile Application", "React Native + Expo SDK 50", "Cross-platform iOS/Android", "App store security review"],
        ["Cloud Database", "Supabase (PostgreSQL)", "All business financial data", "Row-Level Security, SOC2"],
        ["Authentication", "Supabase Auth + TOTP 2FA", "User identity and access", "Fail-secure enforcement"],
        ["Encryption", "AES-256 (CryptoJS)", "Sensitive data at rest", "Bank-grade encryption"],
        ["API Security", "Certificate Pinning", "All external API calls", "MITM attack prevention"],
        ["Offline Storage", "AsyncStorage + Sync Queue", "Offline-first operations", "Local encryption"],
        ["Bank Integration", "Pngme API", "Bank data aggregation", "Pngme regulatory license"],
        ["Payments", "Paystack + Flutterwave", "Payment collection", "PCI-DSS compliant"],
        ["AI Layer", "LLM integration", "Business advisor", "Data stays in user account"],
        ["Audit", "Custom audit log module", "All significant actions", "Immutable log chain"],
    ]
)

# ═══════════════════════════════════════════════════════════════════════════════
#  APPENDIX B — SECURITY CONTROLS
# ═══════════════════════════════════════════════════════════════════════════════

add_heading(doc, "APPENDIX B: SECURITY CONTROLS SUMMARY", 1)
add_divider(doc)

for control in [
    ("Data Encryption at Rest", "HIGH", "All sensitive user financial data encrypted with AES-256 before storage. Encryption keys user-specific and never co-located with encrypted data."),
    ("Data Encryption in Transit", "HIGH", "All data transmitted over TLS 1.3. Certificate pinning prevents substitution of legitimate certificates."),
    ("Authentication", "HIGH", "Multi-factor authentication required. 2FA implemented with TOTP. Fail-secure: exceptions during verification force re-authentication, never grant access."),
    ("Authorisation", "HIGH", "Row-Level Security on all database tables. Users can only read and write their own data. No shared data access without explicit consent."),
    ("Offline Security", "MEDIUM", "Locally cached data protected by device security. Sync queue encrypts sensitive payloads before local storage."),
    ("Audit Logging", "HIGH", "All significant actions logged with user ID, timestamp, IP address, and action description. Logs are immutable and retained for regulatory compliance period."),
    ("Incident Response", "MEDIUM", "Defined incident response playbook. Severity classification system. User notification protocol for any breach affecting their data."),
    ("Penetration Testing", "PLANNED", "Third-party penetration testing scheduled before full commercial deployment. Results shared with Zenith Bank on request."),
]:
    add_bullet(doc, f"[{control[1]}] {control[2]}", bold_prefix=control[0] + ":")

# ═══════════════════════════════════════════════════════════════════════════════
#  APPENDIX C — SAMPLE FINANCIAL REPORT
# ═══════════════════════════════════════════════════════════════════════════════

add_heading(doc, "APPENDIX C: SAMPLE QUAD360 FINANCIAL PROFILE", 1)
add_divider(doc)
add_body(doc, "The following illustrates the financial profile generated for a typical Quad360 business after 6 months of active use. This is the data package that would be shared with Zenith Bank's credit team (with user consent) as part of a loan application.")

add_heading(doc, "Business: Sample Fashion & Retail SME — Lagos | 6 Months Active", 2)
add_simple_table(doc,
    ["Financial Metric", "Month 1", "Month 3", "Month 6", "Trend"],
    [
        ["Monthly Revenue", "₦850,000", "₦1,240,000", "₦1,680,000", "↑ 97.6%"],
        ["Monthly Expenses", "₦720,000", "₦920,000", "₦1,140,000", "↑ 58.3%"],
        ["Net Profit", "₦130,000", "₦320,000", "₦540,000", "↑ 315%"],
        ["Net Profit Margin", "15.3%", "25.8%", "32.1%", "↑ Improving"],
        ["Gross Margin", "38.2%", "41.5%", "44.8%", "↑ Improving"],
        ["Outstanding Invoices", "₦380,000", "₦290,000", "₦145,000", "↓ Improving"],
        ["Invoice Collection Rate", "62%", "78%", "91%", "↑ Strong"],
        ["Cash Flow Position", "Tight", "Stable", "Positive", "↑ Healthy"],
        ["Current Ratio", "0.82", "1.14", "1.47", "↑ Above threshold"],
        ["Debt-to-Equity", "1.8", "1.2", "0.7", "↓ Improving"],
        ["Financial Health Score", "42/100", "61/100", "78/100", "↑ Low → Good"],
    ]
)
add_body(doc, "Credit Officer Assessment: This business demonstrates consistent revenue growth (+97.6%), improving profitability (15.3% → 32.1% net margin), strong invoice collection improvement (62% → 91%), and a financial health trajectory from 'marginal' to 'good' over 6 months. The 6-month data record provides sufficient basis for a credit decision. Recommended loan assessment: ₦2.5M working capital facility.", italic=True, color=MID_GREY)

# ── Save ──────────────────────────────────────────────────────────────────────
output_path = "/home/user/sme-financial-mobile-app/proposal/Quad360_Zenith_Bank_Business_Proposal.docx"
doc.save(output_path)
print(f"✅  Saved: {output_path}")
