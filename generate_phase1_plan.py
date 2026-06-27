#!/usr/bin/env python3
"""
Generate Phase 1 Execution Plan as Microsoft Word Document
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime

def add_heading_with_color(doc, text, level, color_rgb=(0, 51, 102)):
    """Add heading with custom color"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(*color_rgb)
    return heading

def shade_table_header(table, color_rgb=(0, 51, 102)):
    """Shade table header row"""
    for cell in table.rows[0].cells:
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), '%02x%02x%02x' % color_rgb)
        cell._element.get_or_add_tcPr().append(shading_elm)

        # Make text white
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.bold = True

def add_table_with_data(doc, headers, data, header_color=(0, 51, 102)):
    """Add formatted table"""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'

    # Header row
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header

    shade_table_header(table, header_color)

    # Data rows
    for row_data in data:
        row_cells = table.add_row().cells
        for i, cell_data in enumerate(row_data):
            row_cells[i].text = str(cell_data)

    return table

def create_phase1_document():
    """Create comprehensive Phase 1 Execution Plan Word document"""

    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # ============ TITLE PAGE ============
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run('QUAD360')
    title_run.font.size = Pt(36)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0, 51, 102)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run('PHASE 1 EXECUTION PLAN')
    subtitle_run.font.size = Pt(24)
    subtitle_run.font.bold = True
    subtitle_run.font.color.rgb = RGBColor(0, 102, 153)

    tagline = doc.add_paragraph()
    tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tagline_run = tagline.add_run('Building to 3% Market Share in 6 Months')
    tagline_run.font.size = Pt(14)
    tagline_run.font.italic = True
    tagline_run.font.color.rgb = RGBColor(102, 102, 102)

    doc.add_paragraph()  # Spacer

    # Metadata
    meta_table = doc.add_table(rows=6, cols=2)
    meta_table.style = 'Light Grid Accent 1'

    meta_data = [
        ('Period', 'Month 1–6 (Immediate Launch through End-of-Month 6)'),
        ('Target', '1.2M signed-up users, 400k active users, 3% market share'),
        ('Prepared by', 'Quadri Abiodun, Co-Founder & CEO, Quad360'),
        ('Document', 'PHASE_1_EXECUTION_PLAN.md'),
        ('Date', datetime.now().strftime('%B %d, %Y')),
        ('Classification', 'Strategic / Confidential'),
    ]

    for i, (key, value) in enumerate(meta_data):
        meta_table.rows[i].cells[0].text = key
        meta_table.rows[i].cells[1].text = value
        meta_table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True

    doc.add_page_break()

    # ============ TABLE OF CONTENTS ============
    add_heading_with_color(doc, 'TABLE OF CONTENTS', 1)

    toc_items = [
        '1. Executive Summary',
        '2. Critical Path Dependencies',
        '3. Month 1: Credibility + Foundation',
        '4. Month 2: Activation & Community Infiltration',
        '5. Month 3–4: Scale & Ecosystem Solidification',
        '6. Month 5–6: Acceleration & First Milestones',
        '7. Phase 1 Success Criteria',
        '8. Resource Allocation',
        '9. Budget Allocation',
        '10. Contingency Plans',
        '11. Weekly Reporting Cadence',
    ]

    for item in toc_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_page_break()

    # ============ EXECUTIVE SUMMARY ============
    add_heading_with_color(doc, 'Executive Summary', 1)

    doc.add_paragraph(
        'Phase 1 represents the foundation-building stage of Quad360\'s market dominance strategy. '
        'Over 6 months, we will establish credibility through bank partnerships, activate community channels, '
        'and prove viral mechanics. By end of Phase 1, we target 1.2M signups (3% market share) on our path to 60% dominance.'
    )

    key_outcomes = doc.add_paragraph()
    key_outcomes.add_run('Key Outcomes by End of Phase 1:\n').font.bold = True
    outcomes = [
        '✓ 32,000–41,000 cumulative signups',
        '✓ 8,000–12,000 Daily Active Users (DAU)',
        '✓ 500–1,000 paid plan conversions (₦2.5k/month)',
        '✓ 3 viral loops activated (referral, invoices, payment links)',
        '✓ Zenith Bank pilot validated with 5,000+ users',
        '✓ 15 market association partnerships operational',
        '✓ ₦1.25M–₦2.5M Monthly Recurring Revenue (MRR)',
        '✓ >40% Day-30 retention (habit formation)',
        '✓ 50%+ of growth from organic/network effects (not paid acquisition)',
    ]

    for outcome in outcomes:
        doc.add_paragraph(outcome, style='List Bullet')

    doc.add_paragraph()

    investment = doc.add_paragraph()
    investment.add_run('Total Phase 1 Investment: ').font.bold = True
    investment.add_run('₦28M over 6 months\n')
    investment.add_run('Expected CAC (Customer Acquisition Cost): ').font.bold = True
    investment.add_run('₦680–875 per signup\n')
    investment.add_run('Expected LTV (Lifetime Value): ').font.bold = True
    investment.add_run('₦30,000 per paid user (12-month horizon)')

    doc.add_page_break()

    # ============ CRITICAL PATH DEPENDENCIES ============
    add_heading_with_color(doc, 'Critical Path Dependencies', 1)

    doc.add_paragraph(
        'Before any external-facing activity, these foundation items must be operational. '
        'Any delay in these blocks the entire Phase 1 timeline.'
    )

    # Week 1-2 requirements
    add_heading_with_color(doc, 'Week 1–2: Foundation Setup', 2, (0, 102, 153))

    dependencies = [
        {
            'name': 'Zenith Bank Partnership Agreement Signed',
            'owner': 'Quadri (CEO) + Partnerships manager',
            'deliverable': 'Signed pilot agreement committing 500 SMEs, duration 90 days',
            'risk': 'HIGH — if Zenith delays, entire bank distribution plan delays',
            'contingency': 'Parallel track Access Bank conversation'
        },
        {
            'name': 'Platform Stability Audit',
            'owner': 'Engineering lead',
            'deliverable': 'Zero known critical bugs, 99.5% uptime SLA met, security audit complete',
            'risk': 'HIGH — crashes on launch destroy brand reputation',
            'contingency': 'Load test to 10k concurrent users before public announcement'
        },
        {
            'name': 'Customer Support Ticketing System Live',
            'owner': 'Operations manager',
            'deliverable': 'Zendesk/similar ticketing + WhatsApp support channel',
            'risk': 'MEDIUM — slow support = churn',
            'contingency': 'Pre-hire 3 support staff before launch week'
        },
        {
            'name': 'Analytics & KPI Dashboard Live',
            'owner': 'Product + Analytics',
            'deliverable': 'Real-time dashboard: DAU, signups, retention, payments, errors',
            'risk': 'MEDIUM — no visibility = flying blind',
            'contingency': 'Manual Sheets tracking if automated dashboard delayed'
        },
        {
            'name': 'Referral Program Infrastructure',
            'owner': 'Engineering',
            'deliverable': 'Referral codes, tracking, reward distribution automated',
            'risk': 'MEDIUM — broken referral = kills viral growth',
            'contingency': 'Manual referral tracking via spreadsheet (less scalable, but functional)'
        }
    ]

    for dep in dependencies:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f"{dep['name']}\n").font.bold = True
        p.add_run(f"Owner: {dep['owner']}\n")
        p.add_run(f"Deliverable: {dep['deliverable']}\n")
        p.add_run(f"Risk: {dep['risk']}\n").font.color.rgb = RGBColor(204, 0, 0)
        p.add_run(f"Contingency: {dep['contingency']}")

    doc.add_page_break()

    # ============ MONTH 1 ============
    add_heading_with_color(doc, 'MONTH 1: Credibility + Foundation', 1)

    doc.add_paragraph(
        'Month 1 focuses on internal readiness and soft launch validation. We prove the product works, '
        'build support infrastructure, and validate that early users find value.'
    )

    # Week 1 tasks
    add_heading_with_color(doc, 'Week 1: Internal Readiness', 2, (0, 102, 153))

    week1_tasks = [
        ('Zenith Bank Pilot Kickoff', 'Quadri + Partnerships', '1 week', 'Final signed agreement + 500 SME contact list'),
        ('Platform Load Testing', 'Engineering', '3 days', '10k concurrent users, <500ms response time'),
        ('Customer Support Setup', 'Operations', '3 days', 'WhatsApp + Zendesk, 3 trained staff'),
        ('Analytics Dashboard Launch', 'Product', '3 days', 'Real-time KPI tracking for internal team'),
    ]

    table = add_table_with_data(
        doc,
        ['Task', 'Owner', 'Timeline', 'Success Criteria'],
        week1_tasks
    )

    doc.add_paragraph()

    # Week 2 tasks
    add_heading_with_color(doc, 'Week 2: Go-To-Market Preparation', 2, (0, 102, 153))

    week2_tasks = [
        ('Zenith Communications Kit', 'Marketing', '2 days', 'Email + WhatsApp + SMS templates (3 variants)'),
        ('Micro-Influencer Briefing', 'Marketing', '3 days', '30 influencers, briefing doc, access codes'),
        ('Community Partnership Pipeline', 'Growth', '5 days', '10 market associations, MOUs drafted'),
        ('WhatsApp Broadcast Setup', 'Growth', '2 days', '500 contacts, segmented by audience'),
    ]

    table = add_table_with_data(
        doc,
        ['Task', 'Owner', 'Timeline', 'Deliverable'],
        week2_tasks
    )

    doc.add_paragraph()

    # Week 3-4 soft launch
    add_heading_with_color(doc, 'Week 3–4: Soft Launch (500-Person Pilot)', 2, (0, 102, 153))

    soft_launch_tasks = [
        ('Zenith SME Pilot Launch', 'Product + Growth', '1 week', '500 users, 25–30% signup rate'),
        ('Micro-Influencer Content', 'Marketing', '1 week', '30 TikToks + 30 Reels, 2M+ reach'),
        ('Market Association Events', 'Growth', '2 weeks', '3 demo events, 150–300 signups'),
        ('Church Business Groups', 'Growth', 'Ongoing', '200 WhatsApp contacts, 1 broadcast/week'),
    ]

    table = add_table_with_data(
        doc,
        ['Task', 'Owner', 'Timeline', 'Expected Result'],
        soft_launch_tasks
    )

    doc.add_paragraph()

    # Month 1 Success Metrics
    add_heading_with_color(doc, 'Month 1 Success Metrics', 2, (0, 102, 153))

    m1_metrics = [
        ('Zenith pilot MOA signed', '1 ✓'),
        ('Platform stress test passed', 'Yes ✓'),
        ('Support team trained', '3 people'),
        ('Analytics dashboard live', 'Yes ✓'),
        ('Influencer briefing packs sent', '30 influencers'),
        ('Market association MOUs', '3+ signed'),
        ('Cumulative signups', '500–1,000'),
        ('DAU (Daily Active Users)', '150–300'),
        ('D7 retention (% back on day 7)', '>25%'),
    ]

    table = add_table_with_data(
        doc,
        ['Metric', 'Target'],
        m1_metrics
    )

    doc.add_paragraph()

    summary = doc.add_paragraph()
    summary.add_run('Month 1 Expected Outcomes:\n').font.bold = True
    summary.add_run(
        '✓ 500–1,000 cumulative signups\n'
        '✓ 150–300 DAU\n'
        '✓ 0 critical platform issues\n'
        '✓ 3 market association partnerships locked in\n'
        '✓ 30 micro-influencers creating content'
    )

    doc.add_page_break()

    # ============ MONTH 2 ============
    add_heading_with_color(doc, 'MONTH 2: Activation & Community Infiltration', 1)

    doc.add_paragraph(
        'Month 2 shifts from soft launch to full community mobilization. We activate viral loops, '
        'scale community events, and prove organic growth is working.'
    )

    # Key initiatives
    add_heading_with_color(doc, 'Key Initiatives', 2, (0, 102, 153))

    m2_initiatives = [
        ('Market Day Demo Tour', 'Weekly on-the-ground demos at markets (Alaba, Balogun, Computer Village)', '20–30 signups/event'),
        ('Church Fellowship Program', 'Partner with 20 church business fellowships', '240–400 signups per 10 churches'),
        ('Influencer Wave 1 Amplification', '30 influencers post 60 pieces of content', '2M+ impressions, 10k clicks'),
        ('WhatsApp Educational Series', '4-week broadcast series (1x/week)', '25–40 signups per broadcast'),
        ('Referral Program Launch', 'Refer 1 → both get 1 month free', '0.3 k-factor, 600 new users'),
    ]

    table = add_table_with_data(
        doc,
        ['Initiative', 'Description', 'Expected Result'],
        m2_initiatives
    )

    doc.add_paragraph()

    # Month 2 Success Metrics
    add_heading_with_color(doc, 'Month 2 Success Metrics', 2, (0, 102, 153))

    m2_metrics = [
        ('Market day events held', '8'),
        ('Direct signups from market days', '160–320'),
        ('Church partnerships launched', '10'),
        ('Church workshop signups', '200–400'),
        ('Influencer content pieces', '60'),
        ('Influencer content reach', '2M impressions'),
        ('WhatsApp broadcasts sent', '4'),
        ('Referral program active users', '1,500+'),
        ('Month 2 signups', '3,000–4,000'),
        ('Cumulative signups (M1+M2)', '4,000–5,000'),
        ('DAU', '800–1,200'),
        ('D7 retention', '>30%'),
    ]

    table = add_table_with_data(
        doc,
        ['Metric', 'Target'],
        m2_metrics
    )

    doc.add_paragraph()

    summary = doc.add_paragraph()
    summary.add_run('Month 2 Expected Outcomes:\n').font.bold = True
    summary.add_run(
        '✓ 3,000–4,000 new signups (cumulative: 4k–5k)\n'
        '✓ 800–1,200 DAU\n'
        '✓ 3 viral loops activated (referral, invoices, payment links)\n'
        '✓ 10 church partnerships operational\n'
        '✓ 8 market day events completed'
    )

    doc.add_page_break()

    # ============ MONTH 3-4 ============
    add_heading_with_color(doc, 'MONTH 3–4: Scale & Ecosystem Solidification', 1)

    doc.add_paragraph(
        'Month 3–4 focuses on validation of pilot results and transition to paid acquisition. '
        'Bank partnership scales from 500 to 5,000 users. Paid social campaigns launch.'
    )

    add_heading_with_color(doc, 'Zenith Bank Pilot Ramp', 2, (0, 102, 153))

    zenith_ramp = doc.add_paragraph(style='List Bullet')
    zenith_ramp.add_run('Pilot Month 1 Results (Week 9–10): ').font.bold = True
    zenith_ramp.add_run(
        'Measure 500 Zenith SMEs\' engagement: activation rate (target >50%), '
        'D7 retention (target >30%), feature adoption, data quality'
    )

    zenith_ramp2 = doc.add_paragraph(style='List Bullet')
    zenith_ramp2.add_run('Expansion Phase (Week 11–16): ').font.bold = True
    zenith_ramp2.add_run(
        'If pilot succeeds, expand from 500 to 5,000 Zenith SME users via email + in-app announcements'
    )

    doc.add_paragraph()

    add_heading_with_color(doc, 'Broader Community Expansion', 2, (0, 102, 153))

    expansion_data = [
        ('Market Association Coalition', 'Expand from 3 to 15 associations', '3k potential users'),
        ('Paid Social Campaign', 'Instagram + Facebook, ₦500k budget', '2,500–3,000 signups'),
        ('Influencer Wave 2', '30 additional micro-influencers', '2M+ impressions, 3–4k signups'),
        ('Ecosystem Integration Prep', 'FIRS tax filing + Payroll beta test', 'Ready for Phase 2'),
    ]

    table = add_table_with_data(
        doc,
        ['Initiative', 'Target', 'Expected Result'],
        expansion_data
    )

    doc.add_paragraph()

    # Month 3-4 Success Metrics
    add_heading_with_color(doc, 'Month 3–4 Success Metrics', 2, (0, 102, 153))

    m34_metrics = [
        ('Zenith pilot activation rate', '>50%'),
        ('Zenith pilot D7 retention', '>30%'),
        ('Zenith expansion to', '5,000 users'),
        ('Market association partnerships', '15 total'),
        ('Paid social signups/month', '2,500–3,000'),
        ('Influencer Wave 2 reach', '2M+ impressions'),
        ('Month 3–4 signups combined', '8,000–12,000'),
        ('Cumulative signups (M1–M4)', '12,000–17,000'),
        ('DAU', '3,000–5,000'),
        ('Paid plan conversion rate', '2–3% of DAU'),
        ('D7 retention', '>35%'),
    ]

    table = add_table_with_data(
        doc,
        ['Metric', 'Target'],
        m34_metrics
    )

    doc.add_page_break()

    # ============ MONTH 5-6 ============
    add_heading_with_color(doc, 'MONTH 5–6: Acceleration & First Milestones', 1)

    doc.add_paragraph(
        'Month 5–6 focuses on viral loop maturation and first revenue generation. '
        'Organic growth should exceed paid acquisition. Paid tier conversions accelerate.'
    )

    add_heading_with_color(doc, 'Viral Loop Maturation', 2, (0, 102, 153))

    viral_loops = [
        ('Referral Program Analysis', '0.3–0.5 k-factor, 4–5k new users from referral'),
        ('Invoice Loop Optimization', '5–10% recipients download, 20% convert to signup'),
        ('Payment Link Loop', '10–20% of payers explore app, 2–3k signups'),
        ('Network Effects Dashboard', 'Organic growth should be 50%+ of total signups'),
    ]

    table = add_table_with_data(
        doc,
        ['Loop', 'Expected Result'],
        viral_loops
    )

    doc.add_paragraph()

    add_heading_with_color(doc, 'Paid Acquisition Scale', 2, (0, 102, 153))

    doc.add_paragraph('Budget increase from ₦500k to ₦2M/month based on proven CPA <₦50k with ROAS >2x')

    paid_acq = [
        ('Platform', 'Budget', 'Expected Signups/Month'),
        ('Instagram', '₦1M', '2,000–2,500'),
        ('Facebook', '₦500k', '1,000–1,500'),
        ('TikTok (new)', '₦500k', '1,000–1,500'),
        ('Total', '₦2M', '4,000–5,000'),
    ]

    table = add_table_with_data(doc, paid_acq[0], paid_acq[1:])

    doc.add_paragraph()

    add_heading_with_color(doc, 'Product Completeness', 2, (0, 102, 153))

    features = [
        ('Feature', 'Timeline', 'Expected Impact'),
        ('Payroll Feature General Release', 'Week 19', '15% of users with staff enable (1,000+)'),
        ('Cash Flow Forecasting Enhancement', 'Week 21', 'Increased engagement, better stickiness'),
    ]

    table = add_table_with_data(doc, features[0], features[1:])

    doc.add_paragraph()

    # Month 5-6 Success Metrics
    add_heading_with_color(doc, 'Month 5–6 Success Metrics', 2, (0, 102, 153))

    m56_metrics = [
        ('Referral program new users', '4,000–5,000'),
        ('Invoice loop new users', '1,000–2,000'),
        ('Organic growth % of total', '>50%'),
        ('Paid social spend/month', '₦2M'),
        ('Paid social signups/month', '4,000–5,000'),
        ('Official Quad360 social followers', '10k–20k'),
        ('Payroll feature active users', '1,000+'),
        ('Month 5–6 combined signups', '10,000–12,000'),
        ('CUMULATIVE SIGNUPS (M1–M6)', '32,000–41,000'),
        ('DAU (end of M6)', '8,000–12,000'),
        ('Paid conversions cumulative', '500–1,000 users'),
        ('MRR (end of M6)', '₦1.25M–₦2.5M'),
        ('D30 retention', '>40%'),
    ]

    table = add_table_with_data(
        doc,
        ['Metric', 'Target'],
        m56_metrics
    )

    doc.add_page_break()

    # ============ PHASE 1 SUCCESS CRITERIA ============
    add_heading_with_color(doc, 'Phase 1 Success Criteria (End of Month 6)', 1)

    add_heading_with_color(doc, 'User Metrics', 2, (0, 102, 153))

    doc.add_paragraph('32,000–41,000 cumulative signups (on track to 1.2M by end of Year 1)', style='List Bullet')
    doc.add_paragraph('8,000–12,000 DAU (target: 25–30% of signups)', style='List Bullet')
    doc.add_paragraph('>40% D30 retention (users finding value and returning)', style='List Bullet')
    doc.add_paragraph('500–1,000 paid conversions (2–3% of DAU)', style='List Bullet')
    doc.add_paragraph('₦1.25M–₦2.5M MRR (enough to cover costs + reinvest)', style='List Bullet')

    add_heading_with_color(doc, 'Engagement Metrics', 2, (0, 102, 153))

    doc.add_paragraph('3+ daily transactions per DAU (frequency = habit formation)', style='List Bullet')
    doc.add_paragraph('5+ invoices created per active SME per month', style='List Bullet')
    doc.add_paragraph('Payroll adoption: 10%+ of SMEs with staff using feature', style='List Bullet')

    add_heading_with_color(doc, 'Viral Metrics', 2, (0, 102, 153))

    doc.add_paragraph('Referral k-factor: >0.3 (organic growth loop working)', style='List Bullet')
    doc.add_paragraph('Invoice loop conversion: 5%+ of recipients → signup', style='List Bullet')
    doc.add_paragraph('Payment link loop: 2,000+ signups', style='List Bullet')
    doc.add_paragraph('Organic growth: 50%+ of new signups from network effects', style='List Bullet')

    add_heading_with_color(doc, 'Partnership Metrics', 2, (0, 102, 153))

    doc.add_paragraph('Zenith Bank: 5,000+ SMEs using Quad360 via bank channel', style='List Bullet')
    doc.add_paragraph('Market associations: 15+ partnerships operational', style='List Bullet')
    doc.add_paragraph('Influencers: 60 pieces of content, 5M+ cumulative impressions', style='List Bullet')
    doc.add_paragraph('Community: 1,000+ direct signups from in-person market days', style='List Bullet')

    add_heading_with_color(doc, 'Platform Metrics', 2, (0, 102, 153))

    doc.add_paragraph('Uptime: 99.5%+ (zero major outages)', style='List Bullet')
    doc.add_paragraph('Error rate: <0.5% of transactions', style='List Bullet')
    doc.add_paragraph('Data integrity: 100% transaction accuracy on reconciliation', style='List Bullet')
    doc.add_paragraph('Security: Passed third-party penetration test', style='List Bullet')

    add_heading_with_color(doc, 'Financial Health', 2, (0, 102, 153))

    doc.add_paragraph('CAC (Customer Acquisition Cost): <₦50k', style='List Bullet')
    doc.add_paragraph('LTV (Lifetime Value): ₦30k for monthly users', style='List Bullet')
    doc.add_paragraph('Payback period: <12 months', style='List Bullet')
    doc.add_paragraph('Unit economics: Positive for paid tier users', style='List Bullet')

    doc.add_page_break()

    # ============ RESOURCE ALLOCATION ============
    add_heading_with_color(doc, 'Resource Allocation (Phase 1)', 1)

    add_heading_with_color(doc, 'Team Structure', 2, (0, 102, 153))

    team_roles = [
        ('Role', 'Responsibility', 'Headcount'),
        ('CEO/Founder', 'Zenith partnership, board updates, strategic decisions', '1'),
        ('Product Lead', 'Feature prioritization, roadmap, payroll/tax prep', '1'),
        ('Growth Lead', 'Community partnerships, market events, influencers, paid social', '1'),
        ('Marketing Lead', 'Content creation, social media, messaging, PR', '1'),
        ('Engineering Lead', 'Platform stability, integrations, performance', '2'),
        ('Operations Lead', 'Support ticketing, SLA management, partner onboarding', '1'),
        ('Data/Analytics', 'Real-time dashboards, KPI tracking, weekly reports', '1'),
    ]

    table = add_table_with_data(doc, team_roles[0], team_roles[1:])

    doc.add_paragraph()

    team_note = doc.add_paragraph()
    team_note.add_run('Total headcount for Phase 1: ').font.bold = True
    team_note.add_run('8–10 people')

    doc.add_page_break()

    # ============ BUDGET ALLOCATION ============
    add_heading_with_color(doc, 'Budget Allocation (6 Months)', 1)

    budget_data = [
        ('Category', 'Monthly Avg', 'Total 6-Month', 'Notes'),
        ('Personnel (Salaries)', '₦3.3M', '₦20M', 'Competitive rates for Lagos startup'),
        ('Paid Acquisition', '₦1M', '₦3M', 'Ramps from ₦500k to ₦2M by M5'),
        ('Events/Market Days', '₦333k', '₦2M', '8 events @ ₦250k each'),
        ('Content Creation', '₦167k', '₦1M', 'Influencer stipends, video production'),
        ('Infrastructure/Tools', '₦167k', '₦1M', 'Analytics, CRM, SMS broadcast, hosting'),
        ('Contingency', '₦167k', '₦1M', '5% buffer for unknowns'),
    ]

    table = add_table_with_data(doc, budget_data[0], budget_data[1:])

    doc.add_paragraph()

    budget_summary = doc.add_paragraph()
    budget_summary.add_run('TOTAL PHASE 1 BUDGET: ').font.bold = True
    budget_summary.add_run('₦28M over 6 months (₦4.67M/month average)')

    doc.add_paragraph()

    roi = doc.add_paragraph()
    roi.add_run('Expected Return on Investment:\n').font.bold = True
    roi.add_run(
        'Phase 1 Spend: ₦28M\n'
        'Signups Acquired: 32k–41k\n'
        'CAC per Signup: ₦680–875\n'
        'Free → Paid Conversion (3%): 1,000 users\n'
        'MRR at Month 6: ₦1.25M–₦2.5M\n'
        'MRR is 65% of monthly burn by Month 6\n'
        '→ By Month 10–12, MRR covers all costs + creates profit'
    )

    doc.add_page_break()

    # ============ CONTINGENCY PLANS ============
    add_heading_with_color(doc, 'Contingency Plans', 1)

    risks = [
        {
            'name': 'Risk 1: Zenith Bank Pilot Fails (Low Activation)',
            'trigger': 'Zenith pilot activation <30%',
            'response': 'Root cause analysis immediately',
            'action': 'Pivot to Access Bank (parallel track)',
            'timeline': '1 week analysis, 2 weeks course-correct'
        },
        {
            'name': 'Risk 2: Paid Acquisition CPA Too High',
            'trigger': 'CPA >₦100k (unit economics negative)',
            'response': 'Pause paid spending',
            'action': 'Double down on organic (referral + community)',
            'timeline': 'Weeks 7–12, evaluate by Week 12'
        },
        {
            'name': 'Risk 3: Platform Stability Issues',
            'trigger': 'Uptime <99%, error rate >2%',
            'response': 'Halt all marketing immediately',
            'action': 'Go into stability mode, fix critical issues',
            'timeline': 'Immediate, target 48 hours recovery'
        },
        {
            'name': 'Risk 4: Influencer Content Performs Poorly',
            'trigger': 'Influencer CTR <2% (below 5–8% expected)',
            'response': 'Wrong audience or messaging',
            'action': 'Switch to different influencer cohort (business-focused)',
            'timeline': 'Week 7, new wave by Week 9'
        },
        {
            'name': 'Risk 5: Market Day Events Have Low Turnout',
            'trigger': '<50 attendees per event',
            'response': 'Venue or promotion problem',
            'action': 'Partner with association leader for co-hosting',
            'timeline': 'Week 5, all events >100 attendees by Week 7'
        }
    ]

    for risk in risks:
        add_heading_with_color(doc, risk['name'], 2, (153, 0, 0))

        p = doc.add_paragraph()
        p.add_run('Trigger: ').font.bold = True
        p.add_run(risk['trigger'] + '\n')

        p.add_run('Response: ').font.bold = True
        p.add_run(risk['response'] + '\n')

        p.add_run('Action: ').font.bold = True
        p.add_run(risk['action'] + '\n')

        p.add_run('Timeline: ').font.bold = True
        p.add_run(risk['timeline'])

        doc.add_paragraph()

    doc.add_page_break()

    # ============ WEEKLY REPORTING ============
    add_heading_with_color(doc, 'Weekly Reporting Cadence', 1)

    doc.add_paragraph(
        'Every Friday 5pm, CEO + Growth Lead + Product Lead review these metrics:'
    )

    reporting = [
        'Signups This Week (actual vs. target)',
        'DAU (current levels, trend)',
        'Retention (D7, D30 tracking)',
        'Top Signups Channel (where are users coming from?)',
        'Top Churn Reason (exit surveys, support tickets)',
        'Revenue (MRR, paid conversions)',
        'Next Week Priorities (3 biggest bets)',
        'Blockers (what needs to be unblocked?)',
    ]

    for item in reporting:
        doc.add_paragraph(item, style='List Number')

    doc.add_paragraph()

    escalation = doc.add_paragraph()
    escalation.add_run('Escalation Protocol: ').font.bold = True
    escalation.add_run(
        'If any metric is >20% below target, escalate to full team + board for emergency response.'
    )

    doc.add_page_break()

    # ============ CLOSING ============
    add_heading_with_color(doc, 'Success Looks Like (End of Phase 1)', 1)

    success_items = [
        '✅ 1.2M signups (on track to hit)',
        '✅ 40%+ D30 retention (users finding value)',
        '✅ Viral loops working (referral k-factor >0.3, organic growth >50%)',
        '✅ Bank partnership locked in (5k users from Zenith, path to scale)',
        '✅ Community infrastructure in place (15 associations, weekly market presence)',
        '✅ Paid acquisition proven (CPA <₦50k, positive unit economics)',
        '✅ Revenue baseline (₦1.25M–₦2.5M MRR, covers costs + reinvest)',
        '✅ Platform 100% stable (99.5%+ uptime, zero critical incidents)',
    ]

    for item in success_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph()

    closing = doc.add_paragraph()
    closing.add_run(
        'This plan is NOT aspirational — every number is based on comparable SaaS benchmarks, '
        'Zenith partnership data, and market research. Execute with discipline. Measure weekly. '
        'Course-correct immediately.'
    ).font.italic = True

    doc.add_paragraph()

    next = doc.add_paragraph()
    next.add_run('Next milestone: ').font.bold = True
    next.add_run('Week 24 (End of Phase 1) → Phase 2 Kickoff (Months 7–18)')

    doc.add_page_break()

    # ============ FOOTER ============
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run('QUAD360 — Building Financial Operating Systems for African SMEs').font.italic = True
    footer.add_run('\n')
    footer.add_run(f'Document Generated: {datetime.now().strftime("%B %d, %Y")}\n')
    footer.add_run('Confidential & Strategic Use Only')

    # Save document
    doc.save('/home/user/sme-financial-mobile-app/QUAD360_PHASE_1_EXECUTION_PLAN.docx')
    print('✅ Document created: QUAD360_PHASE_1_EXECUTION_PLAN.docx')

if __name__ == '__main__':
    create_phase1_document()
